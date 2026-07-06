from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from .model import Finding, ReportModel
from .styles import result_color


def _hex_to_rgb(h: str) -> RGBColor:
    h = h.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _add_finding(doc: Document, f: Finding) -> None:
    head = doc.add_paragraph()
    badge = head.add_run(f"[{f.result_label}] ")
    badge.bold = True
    badge.font.color.rgb = _hex_to_rgb(result_color(f.result_key))
    head.add_run(f"{f.rule_code} {f.name}")
    if f.severity is not None:
        head.add_run(f"    严重度：{f.severity}")
    if f.confidence is not None:
        doc.add_paragraph(f"置信度：{f.confidence:.2f}")
    doc.add_paragraph(f"审查意见：{f.suggestion or '—'}")
    doc.add_paragraph("依据出处：")
    if f.evidence:
        for e in f.evidence:
            doc.add_paragraph(f"· {e.locator}：{e.quote}")
    else:
        doc.add_paragraph("无")
    if f.audit is not None:
        doc.add_paragraph("—— 审计留痕 ——")
        doc.add_paragraph(f"机审初判：{f.audit.initial} → 人工改判：{f.audit.final}")
        doc.add_paragraph(f"处置说明：{f.audit.disposition or '—'}")


def render_docx(m: ReportModel) -> bytes:
    doc = Document()

    # ① 封面
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(m.title)
    tr.bold = True
    tr.font.size = Pt(24)
    doc.add_paragraph()
    for label, value in m.cover:
        doc.add_paragraph(f"{label}：{value}")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run(m.footer_note)
    doc.add_page_break()

    # ② 结论页
    doc.add_heading("一、审查结论", level=1)
    doc.add_paragraph(m.conclusion_text)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "维度", "通过", "不通过", "需关注"
    for st in m.dimension_stats:
        row = tbl.add_row().cells
        row[0].text = st.dimension_label
        row[1].text = str(st.passed)
        row[2].text = str(st.failed)
        row[3].text = str(st.attention)

    # ③ 分维度发现
    doc.add_heading("二、分维度审查发现", level=1)
    for sec in m.sections:
        doc.add_heading(sec.dimension_label, level=2)
        for f in sec.findings:
            _add_finding(doc, f)

    # ④ 结尾签署
    doc.add_heading("三、审查说明", level=1)
    doc.add_paragraph("本报告由智能立项审查系统依据规则库自动生成，人工复核记录见各条审计留痕。")
    doc.add_paragraph("审查人：______________        日期：______________")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
