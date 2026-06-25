from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

import pdfplumber
from docx import Document
from sqlalchemy import delete, select
from sqlmodel import Session

from . import ocr  # OCR 可选模块；通过模块属性访问以便测试 monkeypatch
from .models import FileObject, ParseSegment, StandardDoc
from .schemas import RecognizeResult
from .storage import FileStorage

MAX_SEGMENT_CHARS = 65536

# DB 约束：segment_type IN ('text','table','title','figure')
_SEGMENT_TYPE_MAP = {"heading": "title", "paragraph": "text"}
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"}


@dataclass
class SegmentDraft:
    page_no: int | None
    locator: dict
    segment_type: str  # heading | paragraph | table
    content_text: str


class UnsupportedFormatError(Exception):
    def __init__(self, ext: str):
        self.ext = ext
        super().__init__(f"unsupported format: {ext}")


def _clip(text: str) -> str:
    return text if len(text) <= MAX_SEGMENT_CHARS else text[:MAX_SEGMENT_CHARS]


def _rows_to_text(rows) -> str:
    lines = [" | ".join((c or "").strip() for c in row) for row in rows]
    return "\n".join(lines).strip()


def parse_pdf(path: Path) -> list[SegmentDraft]:
    drafts: list[SegmentDraft] = []
    fdoc = None
    if ocr.OCR_AVAILABLE:
        import fitz
        fdoc = fitz.open(str(path))
    try:
        with pdfplumber.open(str(path)) as pdf:
            for pageno, page in enumerate(pdf.pages, start=1):
                tables = page.find_tables()
                text_region = page
                for t in tables:
                    text_region = text_region.outside_bbox(t.bbox)
                text = text_region.extract_text() or ""
                blocks = [b.strip() for b in re.split(r"\n\s*\n", text) if b.strip()]
                if not blocks and not tables and fdoc is not None:
                    # 该页无文本层，疑似扫描页 -> OCR 回退
                    lines = ocr.ocr_pdf_page(fdoc[pageno - 1])
                    ocr_text = "\n".join(lines).strip()
                    if ocr_text:
                        drafts.append(SegmentDraft(
                            pageno, {"page": pageno, "block_index": 0, "ocr": True},
                            "paragraph", _clip(ocr_text)))
                    continue
                for bi, block in enumerate(blocks):
                    drafts.append(SegmentDraft(pageno, {"page": pageno, "block_index": bi}, "paragraph", _clip(block)))
                for ti, t in enumerate(tables):
                    ttext = _rows_to_text(t.extract())
                    if ttext:
                        drafts.append(SegmentDraft(pageno, {"page": pageno, "table_index": ti}, "table", _clip(ttext)))
    finally:
        if fdoc is not None:
            fdoc.close()
    return drafts


def _heading_level(style_name: str) -> int | None:
    parts = style_name.split()
    if parts and parts[-1].isdigit():
        return int(parts[-1])
    return None


def parse_docx(path: Path) -> list[SegmentDraft]:
    doc = Document(str(path))
    drafts: list[SegmentDraft] = []
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if not txt:
            continue
        style = (p.style.name if p.style else "") or ""
        if style.startswith("Heading"):
            drafts.append(SegmentDraft(None, {"para_index": i, "level": _heading_level(style)}, "heading", _clip(txt)))
        else:
            drafts.append(SegmentDraft(None, {"para_index": i}, "paragraph", _clip(txt)))
    for j, table in enumerate(doc.tables):
        rows = [[c.text for c in row.cells] for row in table.rows]
        ttext = _rows_to_text(rows)
        if ttext:
            drafts.append(SegmentDraft(None, {"table_index": j}, "table", _clip(ttext)))
    return drafts


def parse_image(path: Path) -> list[SegmentDraft]:
    if not ocr.OCR_AVAILABLE:
        return []
    lines = ocr.ocr_image(path)
    text = "\n".join(lines).strip()
    if not text:
        return []
    return [SegmentDraft(1, {"page": 1, "block_index": 0, "ocr": True}, "paragraph", _clip(text))]


def parse_file(path: Path, ext: str) -> tuple[list[SegmentDraft], str | None]:
    """按扩展名解析为 SegmentDraft 列表。返回 (drafts, error)；error 非空表示无产出原因。"""
    drafts: list[SegmentDraft] = []
    error: str | None = None
    try:
        if not path.exists():
            return [], "源文件缺失，无法识别"
        if ext == ".pdf":
            drafts = parse_pdf(path)
            if not drafts:
                error = ("未从 PDF 抽取到文本，可能是扫描件且 OCR 未识别出文字" if ocr.OCR_AVAILABLE
                         else "未从 PDF 抽取到文本，疑似扫描件；未安装 OCR 组件，请 pip install rapidocr-onnxruntime pymupdf 后重试")
        elif ext == ".docx":
            drafts = parse_docx(path)
            if not drafts:
                error = "文档为空，未抽取到任何内容"
        elif ext in IMAGE_EXTS:
            drafts = parse_image(path)
            if not drafts:
                error = ("未从图片中识别到文字" if ocr.OCR_AVAILABLE
                         else "检测到图片文件，但未安装 OCR 组件，请 pip install rapidocr-onnxruntime pymupdf 后重试")
        else:
            raise UnsupportedFormatError(ext or "(无扩展名)")
    except UnsupportedFormatError as e:
        return [], f"不支持的格式 {e.ext}，请转存为 .docx 或 PDF"
    except Exception as e:  # noqa: BLE001
        return [], f"识别失败：{str(e)[:300]}"
    return drafts, error


def recognize_standard_doc(db: Session, storage: FileStorage, doc_id: int) -> RecognizeResult:
    sd = db.get(StandardDoc, doc_id)
    if sd is None:
        return RecognizeResult(
            doc_id=doc_id, doc_code="", recognition_status="failed",
            segment_count=0, page_count=None, error="文档记录不存在",
        )
    fo = db.get(FileObject, sd.file_id) if sd.file_id else None
    if fo is None:
        drafts, error = [], "源文件记录缺失，无法识别"
    else:
        ext = Path(fo.file_name).suffix.lower()
        drafts, error = parse_file(storage.base_dir / fo.object_key, ext)
    status = "failed"

    if drafts and error is None:
        db.execute(delete(ParseSegment).where(ParseSegment.standard_doc_id == doc_id))
        for d in drafts:
            db.add(ParseSegment(
                standard_doc_id=doc_id,
                material_file_id=None,
                page_no=d.page_no,
                locator=d.locator,
                segment_type=_SEGMENT_TYPE_MAP.get(d.segment_type, d.segment_type),
                content_text=d.content_text,
            ))
        status = "done"
        segment_count = len(drafts)
        pages = {d.page_no for d in drafts if d.page_no is not None}
        page_count = max(pages) if pages else None
    else:
        status, segment_count, page_count = "failed", 0, None
        if error is None:
            error = "识别失败：未产出任何片段"

    sd.recognition_status = status
    sd.recognition_error = error if status == "failed" else None
    db.add(sd)
    db.commit()
    return RecognizeResult(
        doc_id=doc_id, doc_code=sd.doc_code, recognition_status=status,
        segment_count=segment_count, page_count=page_count, error=error if status == "failed" else None,
    )


def recognize_material_file(db: Session, storage: FileStorage, material_file_id: int) -> "MaterialRecognizeResult":
    from .models import MaterialFile
    from .schemas import MaterialRecognizeResult

    mf = db.get(MaterialFile, material_file_id)
    if mf is None:
        return MaterialRecognizeResult(material_file_id=material_file_id, recognition_status="failed",
                                       segment_count=0, page_count=None, error="材料记录不存在")
    fo = db.get(FileObject, mf.file_id) if mf.file_id else None
    if fo is None:
        drafts, error = [], "源文件记录缺失，无法识别"
    else:
        ext = Path(fo.file_name).suffix.lower()
        drafts, error = parse_file(storage.base_dir / fo.object_key, ext)

    if drafts and error is None:
        db.execute(delete(ParseSegment).where(ParseSegment.material_file_id == material_file_id))
        for d in drafts:
            db.add(ParseSegment(
                material_file_id=material_file_id, standard_doc_id=None,
                page_no=d.page_no, locator=d.locator,
                segment_type=_SEGMENT_TYPE_MAP.get(d.segment_type, d.segment_type),
                content_text=d.content_text,
            ))
        status, segment_count = "done", len(drafts)
        pages = {d.page_no for d in drafts if d.page_no is not None}
        page_count = max(pages) if pages else None
    else:
        status, segment_count, page_count = "failed", 0, None
        if error is None:
            error = "识别失败：未产出任何片段"

    mf.recognition_status = status
    mf.recognition_error = error if status == "failed" else None
    db.add(mf)
    db.commit()
    return MaterialRecognizeResult(
        material_file_id=material_file_id, recognition_status=status,
        segment_count=segment_count, page_count=page_count,
        error=error if status == "failed" else None,
    )


def reset_stuck_processing(db: Session) -> int:
    """应用启动时把残留 processing(进程内后台任务因重启中断) 重置为 pending。"""
    rows = db.execute(select(StandardDoc).where(StandardDoc.recognition_status == "processing")).scalars().all()
    for sd in rows:
        sd.recognition_status = "pending"
        db.add(sd)
    from .models import MaterialFile
    mrows = db.execute(select(MaterialFile).where(MaterialFile.recognition_status == "processing")).scalars().all()
    for mf in mrows:
        mf.recognition_status = "pending"
        db.add(mf)
    db.commit()
    return len(rows) + len(mrows)
