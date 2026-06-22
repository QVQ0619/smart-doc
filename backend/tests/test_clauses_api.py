import io

from docx import Document


def _upload_docx(client, name="规范.docx"):
    d = Document()
    d.add_heading("第一章", level=1)
    d.add_paragraph("第一条 申请人应当具有高级职称。")
    buf = io.BytesIO()
    d.save(buf)
    return client.post(
        "/api/standard-docs",
        files=[("files", (name, buf.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))],
    ).json()["uploaded"][0]


def test_get_segments_returns_doc_segments(client):
    doc = _upload_docx(client)
    segs = client.get(f"/api/standard-docs/{doc['id']}/segments").json()
    assert len(segs) >= 1
    assert all("content_text" in s and "segment_type" in s and "id" in s for s in segs)


def test_post_and_get_clauses_with_provenance(client):
    doc = _upload_docx(client)
    segs = client.get(f"/api/standard-docs/{doc['id']}/segments").json()
    seg_id = segs[0]["id"]
    r = client.post(f"/api/standard-docs/{doc['id']}/clauses", json={"clauses": [
        {"clause_no": "第一条", "clause_text": "申请人应当具有高级职称。", "source_segment_id": seg_id},
        {"clause_no": "#2", "clause_text": "x", "source_segment_id": 999999},
    ]})
    assert r.status_code == 200
    body = r.json()
    assert body["inserted"] == 2
    assert body["missing_provenance"] == 1
    clauses = client.get(f"/api/standard-docs/{doc['id']}/clauses").json()
    assert len(clauses) == 2
    first = next(c for c in clauses if c["clause_no"] == "第一条")
    assert first["source_segment_id"] == seg_id
    assert first["page_no"] == segs[0]["page_no"]  # 出处来自 join parse_segment（docx 为 None）
    bad = next(c for c in clauses if c["clause_no"] == "#2")
    assert bad["source_segment_id"] is None and bad["page_no"] is None


def test_clauses_idempotent_replace(client):
    doc = _upload_docx(client)
    sid = client.get(f"/api/standard-docs/{doc['id']}/segments").json()[0]["id"]
    payload = {"clauses": [{"clause_no": "第一条", "source_segment_id": sid}]}
    client.post(f"/api/standard-docs/{doc['id']}/clauses", json=payload)
    client.post(f"/api/standard-docs/{doc['id']}/clauses", json=payload)
    assert len(client.get(f"/api/standard-docs/{doc['id']}/clauses").json()) == 1


def test_post_empty_clauses_ok(client):
    doc = _upload_docx(client)
    r = client.post(f"/api/standard-docs/{doc['id']}/clauses", json={"clauses": []})
    assert r.status_code == 200 and r.json()["inserted"] == 0


def test_unknown_doc_404(client):
    assert client.get("/api/standard-docs/999999/segments").status_code == 404
    assert client.post("/api/standard-docs/999999/clauses", json={"clauses": []}).status_code == 404
    assert client.get("/api/standard-docs/999999/clauses").status_code == 404
