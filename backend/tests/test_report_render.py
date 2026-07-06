import io
import zipfile
import pdfplumber
from docx import Document
from app.reporting import model, styles
from app.reporting.render_docx import render_docx
from app.reporting.render_pdf import render_pdf
from app.reporting.package import package_zip


def test_report_model_constructs_and_holds_fields():
    ev = model.EvidenceRef(quote="申请人:张三", locator="申请书/第1页")
    audit = model.AuditTrail(initial="通过", final="不通过", disposition="预算表缺失")
    f = model.Finding(result_label="不通过", result_key="fail", rule_code="R-1",
                      name="必须有预算表", severity=3, confidence=0.9,
                      suggestion="应补充预算表", evidence=[ev], audit=audit)
    sec = model.Section(dimension_label="完整性", findings=[f])
    stat = model.DimensionStat(dimension_label="完整性", passed=1, failed=1, attention=0)
    rm = model.ReportModel(title="立项审查报告", cover=[("项目名称", "X")],
                           conclusion_text="综合结论：需整改", dimension_stats=[stat],
                           sections=[sec], footer_note="系统生成")
    assert rm.title == "立项审查报告"
    assert rm.sections[0].findings[0].audit.final == "不通过"
    assert rm.cover[0] == ("项目名称", "X")


def test_styles_dimension_order_and_labels_match_backend():
    assert styles.DIMENSION_ORDER == ["completeness", "normativeness", "compliance",
                                      "consistency", "rationality", "authenticity"]
    assert styles.DIMENSION_LABELS["completeness"] == "完整性"
    assert styles.RESULT_LABELS["fail"] == "不通过"
    assert styles.CONCLUSION_LABELS["fix"] == "需整改"


def test_result_bucket_three_way():
    assert styles.result_bucket("pass") == "passed"
    assert styles.result_bucket("fail") == "failed"
    assert styles.result_bucket("need_review") == "attention"
    assert styles.result_bucket("not_applicable") == "attention"


def _sample_model():
    ev = model.EvidenceRef(quote="申请人:张三", locator="申请书/第1页")
    audit = model.AuditTrail(initial="通过", final="不通过", disposition="预算表缺失")
    f_over = model.Finding(result_label="不通过", result_key="fail", rule_code="R-012",
                           name="必须有预算表", severity=3, confidence=0.92,
                           suggestion="未提供预算明细表，应补充。", evidence=[ev], audit=audit)
    f_plain = model.Finding(result_label="通过", result_key="pass", rule_code="R-003",
                            name="项目名称完整", severity=None, confidence=None,
                            suggestion="", evidence=[], audit=None)
    sec = model.Section(dimension_label="完整性", findings=[f_over, f_plain])
    stat = model.DimensionStat(dimension_label="完整性", passed=1, failed=1, attention=0)
    return model.ReportModel(
        title="立项审查报告",
        cover=[("项目名称", "示范科研项目"), ("申报单位", "某研究所"), ("审查结论", "需整改")],
        conclusion_text="共审查规则 2 条，通过 1、不通过 1、需关注 0。综合结论：需整改",
        dimension_stats=[stat], sections=[sec], footer_note="智能立项审查系统 生成")


def test_render_docx_returns_readable_docx_with_key_text():
    data = render_docx(_sample_model())
    assert isinstance(data, bytes) and len(data) > 0
    doc = Document(io.BytesIO(data))
    text = "\n".join(p.text for p in doc.paragraphs)
    table_text = "\n".join(c.text for t in doc.tables for row in t.rows for c in row.cells)
    whole = text + "\n" + table_text
    assert "立项审查报告" in whole          # 标题
    assert "示范科研项目" in whole          # 封面字段
    assert "综合结论：需整改" in whole       # 结论段
    assert "完整性" in whole                # 维度标题
    assert "必须有预算表" in whole          # 发现名称
    assert "未提供预算明细表" in whole       # 审查意见
    assert "申请书/第1页" in whole          # 出处定位
    assert "机审初判" in whole and "预算表缺失" in whole   # 审计留痕


def test_render_docx_evidence_line_not_list_bullet_style():
    # 出处行不再用 Word 'List Bullet' 样式（避免与手写 '· ' 叠成双项目符号）
    data = render_docx(_sample_model())
    doc = Document(io.BytesIO(data))
    ev_paras = [p for p in doc.paragraphs if "申请书/第1页" in p.text]
    assert ev_paras, "应存在出处段落"
    for p in ev_paras:
        assert (p.style.name or "") != "List Bullet"
        assert p.text.startswith("· ")   # 仍保留单一手写项目符号


def test_render_pdf_is_pdf_and_contains_chinese_text():
    data = render_pdf(_sample_model())
    assert isinstance(data, bytes) and data[:4] == b"%PDF"
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        text = "\n".join((page.extract_text() or "") for page in pdf.pages)
    assert "立项审查报告" in text      # 标题（中文未乱码）
    assert "必须有预算表" in text      # 发现名称
    assert "机审初判" in text          # 审计留痕


def test_render_pdf_escapes_untrusted_xml_special_chars():
    ev = model.EvidenceRef(quote="预算 A<B & C>D 见附表", locator="申请书/第2页")
    f = model.Finding(result_label="不通过", result_key="fail", rule_code="R-099",
                      name="经费<50万 校验", severity=2, confidence=0.8,
                      suggestion="应 <补充> 明细 & 说明", evidence=[ev], audit=None)
    sec = model.Section(dimension_label="完整性", findings=[f])
    stat = model.DimensionStat(dimension_label="完整性", passed=0, failed=1, attention=0)
    rm = model.ReportModel(
        title="立项审查报告",
        cover=[("申报单位", "A&B<研究>所")],
        conclusion_text="综合结论：需整改",
        dimension_stats=[stat], sections=[sec], footer_note="系统生成")
    data = render_pdf(rm)
    assert isinstance(data, bytes) and data[:4] == b"%PDF"
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        text = "\n".join((page.extract_text() or "") for page in pdf.pages)
    assert "A&B<研究>所" in text
    assert "预算 A<B & C>D 见附表" in text
    assert "应 <补充> 明细 & 说明" in text


def test_package_zip_contains_all_files():
    data = package_zip({"a.docx": b"AAA", "b.pdf": b"%PDF-BB"})
    assert isinstance(data, bytes)
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        names = set(z.namelist())
        assert names == {"a.docx", "b.pdf"}
        assert z.read("a.docx") == b"AAA"
        assert z.read("b.pdf") == b"%PDF-BB"
