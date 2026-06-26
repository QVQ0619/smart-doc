"""T2 复现测试：已抽过规则的 standard_doc 重新识别前须级联清理旧 clause/rule。

Bug 根因：regulation_clause.source_segment_id 外键 RESTRICT 引用 parse_segment，
重新识别删旧段落时报外键冲突。
"""
import uuid
from pathlib import Path

from docx import Document
from sqlalchemy import text as sqltext
from sqlmodel import Session

from app.db import engine
from app.dimensions import ensure_dimensions
from app.models import (FileObject, ParseSegment, RegulationClause,
                        ReviewRule, ReviewRuleClause, ReviewRuleVersion,
                        StandardDoc)
from app.recognition import recognize_standard_doc
from app.storage import FileStorage


def _clear_tables():
    with Session(engine) as s:
        s.execute(sqltext("SET FOREIGN_KEY_CHECKS=0"))
        for t in ("review_rule_clause", "review_rule_version",
                  "review_rule", "regulation_clause",
                  "parse_segment", "standard_doc", "file_object"):
            s.execute(sqltext(f"DELETE FROM {t}"))
        s.execute(sqltext("SET FOREIGN_KEY_CHECKS=1"))
        s.commit()


def _make_docx(path: Path) -> None:
    d = Document()
    d.add_heading("章节标题", level=1)
    d.add_paragraph("正文内容一段。")
    d.save(str(path))


def _seed_with_rules(storage: FileStorage, src: Path) -> tuple[int, int, int]:
    """建 standard_doc + 旧 parse_segment + regulation_clause + review_rule 完整链。
    返回 (doc_id, clause_id, rule_id)。"""
    _clear_tables()
    with Session(engine) as s:
        ensure_dimensions(s)

    object_key = "standard_doc/test_rerecog.docx"
    dest = storage.base_dir / object_key
    dest.parent.mkdir(parents=True, exist_ok=True)
    dest.write_bytes(src.read_bytes())

    with Session(engine) as s:
        fo = FileObject(
            bucket="local", object_key=object_key,
            file_name="test_rerecog.docx",
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            size_bytes=src.stat().st_size,
        )
        s.add(fo)
        s.flush()

        sd = StandardDoc(
            doc_code="SD-rerecog-001", title="测试重识别", version="V1.0",
            file_id=fo.id, is_active=True, recognition_status="done",
        )
        s.add(sd)
        s.flush()
        doc_id = sd.id

        # 旧 parse_segment（重识别会先删它，若有 FK 引用则报错）
        old_seg = ParseSegment(
            standard_doc_id=doc_id, material_file_id=None,
            page_no=None, locator={"para_index": 0},
            segment_type="text", content_text="旧内容段落",
        )
        s.add(old_seg)
        s.flush()
        old_seg_id = old_seg.id

        # regulation_clause 引用旧 segment → 触发 FK RESTRICT 的根因
        clause = RegulationClause(
            standard_doc_id=doc_id, doc_code="SD-rerecog-001",
            clause_no="一(一)", clause_text="旧条款文本",
            source_segment_id=old_seg_id,
        )
        s.add(clause)
        s.flush()
        clause_id = clause.id

        # review_rule 全链：rule → version → rule_clause
        dim_id = s.execute(sqltext("SELECT id FROM review_dimension LIMIT 1")).scalar_one()
        rule = ReviewRule(rule_code="RULE-" + uuid.uuid4().hex[:12], current_version_id=None)
        s.add(rule)
        s.flush()
        rule_id = rule.id

        ver = ReviewRuleVersion(
            rule_id=rule_id, version="V1.0", dimension_id=dim_id,
            name="测试规则", logic=None, decision_type="hard",
            disposition="reject", binding_class="common",
        )
        s.add(ver)
        s.flush()

        rule.current_version_id = ver.id
        s.add(rule)
        s.add(ReviewRuleClause(rule_version_id=ver.id, clause_id=clause_id))
        s.commit()

    return doc_id, clause_id, rule_id


def test_rerecognize_clears_clauses_and_rules(tmp_path, storage_dir, _test_db):
    """已抽过规则的 doc 重新识别：级联清旧 clause/rule 后识别成功，旧数据全部清空。"""
    storage = FileStorage(storage_dir)
    src = tmp_path / "test_rerecog.docx"
    _make_docx(src)

    doc_id, _clause_id, rule_id = _seed_with_rules(storage, src)

    # 前置断言：旧数据已存在
    with Session(engine) as s:
        assert s.execute(
            sqltext("SELECT COUNT(*) FROM regulation_clause WHERE standard_doc_id=:d"),
            {"d": doc_id},
        ).scalar_one() == 1, "seed 应建了 1 条 regulation_clause"
        assert s.execute(
            sqltext("SELECT COUNT(*) FROM review_rule WHERE id=:r"),
            {"r": rule_id},
        ).scalar_one() == 1, "seed 应建了 1 条 review_rule"
        assert s.execute(
            sqltext("SELECT COUNT(*) FROM parse_segment WHERE standard_doc_id=:d"),
            {"d": doc_id},
        ).scalar_one() == 1, "seed 应建了 1 条旧 parse_segment"

    # 执行重新识别（bug 时因外键 RESTRICT 崩溃，返回 failed 或抛异常）
    with Session(engine) as db:
        res = recognize_standard_doc(db, storage, doc_id)

    # 识别必须成功
    assert res.recognition_status == "done", (
        f"期望 recognition_status='done'，实际='{res.recognition_status}'，error={res.error!r}"
    )
    assert res.segment_count >= 1

    # 旧条款、旧规则已全部清空
    with Session(engine) as s:
        clause_count = s.execute(
            sqltext("SELECT COUNT(*) FROM regulation_clause WHERE standard_doc_id=:d"),
            {"d": doc_id},
        ).scalar_one()
        rule_count = s.execute(
            sqltext("SELECT COUNT(*) FROM review_rule WHERE id=:r"),
            {"r": rule_id},
        ).scalar_one()
        new_seg_count = s.execute(
            sqltext("SELECT COUNT(*) FROM parse_segment WHERE standard_doc_id=:d"),
            {"d": doc_id},
        ).scalar_one()

    assert clause_count == 0, f"旧 regulation_clause 未清空，仍剩 {clause_count} 条"
    assert rule_count == 0, f"旧 review_rule 未清空，仍剩 {rule_count} 条"
    assert new_seg_count >= 1, "重新识别后应写入新 parse_segment"
