from pathlib import Path

from docx import Document
from sqlalchemy import text as sqltext
from sqlmodel import Session

from app.db import engine
from app.models import FileObject, ParseSegment, StandardDoc
from app.recognition import recognize_standard_doc
from app.storage import FileStorage


def _seed_doc(storage: FileStorage, src: Path, file_name: str) -> int:
    """把 src 落到 storage 并建 file_object+standard_doc，返回 doc_id。"""
    object_key = f"standard_doc/{file_name}"
    dest = storage.base_dir / object_key
    dest.parent.mkdir(parents=True, exist_ok=True)
    dest.write_bytes(src.read_bytes())
    with Session(engine) as s:
        s.execute(sqltext("SET FOREIGN_KEY_CHECKS=0"))
        s.execute(sqltext("DELETE FROM parse_segment"))
        s.execute(sqltext("DELETE FROM standard_doc"))
        s.execute(sqltext("DELETE FROM file_object"))
        s.execute(sqltext("SET FOREIGN_KEY_CHECKS=1"))
        fo = FileObject(bucket="local", object_key=object_key, file_name=file_name, mime_type="x", size_bytes=1)
        s.add(fo)
        s.flush()
        sd = StandardDoc(doc_code="SD-test001", title="t", version="V1.0", file_id=fo.id, is_active=True)
        s.add(sd)
        s.commit()
        return sd.id


def _make_docx(path: Path) -> None:
    d = Document()
    d.add_heading("第一章", level=1)
    d.add_paragraph("正文一段。")
    d.save(str(path))


def _count_segments(doc_id: int) -> int:
    with Session(engine) as s:
        return s.execute(
            sqltext("SELECT COUNT(*) FROM parse_segment WHERE standard_doc_id=:d"), {"d": doc_id}
        ).scalar_one()


def test_recognize_docx_writes_segments_and_done(tmp_path, storage_dir, _test_db):
    storage = FileStorage(storage_dir)
    src = tmp_path / "r.docx"
    _make_docx(src)
    doc_id = _seed_doc(storage, src, "r.docx")
    with Session(engine) as db:
        res = recognize_standard_doc(db, storage, doc_id)
    assert res.recognition_status == "done"
    assert res.segment_count >= 2
    assert res.page_count is None  # docx 无页码
    assert _count_segments(doc_id) == res.segment_count
    with Session(engine) as s:
        sd = s.get(StandardDoc, doc_id)
        assert sd.recognition_status == "done"
        seg = s.execute(
            sqltext("SELECT standard_doc_id, material_file_id FROM parse_segment WHERE standard_doc_id=:d LIMIT 1"),
            {"d": doc_id},
        ).first()
    assert seg[0] == doc_id and seg[1] is None
    with Session(engine) as s:
        types = s.execute(sqltext(
            "SELECT DISTINCT segment_type FROM parse_segment WHERE standard_doc_id=:d"), {"d": doc_id}
        ).scalars().all()
    assert set(types) <= {"text", "table", "title", "figure"}
    assert types  # 非空


def test_recognize_is_idempotent_replace(tmp_path, storage_dir, _test_db):
    storage = FileStorage(storage_dir)
    src = tmp_path / "r.docx"
    _make_docx(src)
    doc_id = _seed_doc(storage, src, "r.docx")
    with Session(engine) as db:
        first = recognize_standard_doc(db, storage, doc_id).segment_count
    with Session(engine) as db:
        second = recognize_standard_doc(db, storage, doc_id).segment_count
    assert first == second
    assert _count_segments(doc_id) == second  # 替换而非累加


def test_recognize_unsupported_format_fails_without_raising(tmp_path, storage_dir, _test_db):
    storage = FileStorage(storage_dir)
    src = tmp_path / "r.doc"
    src.write_bytes(b"\xd0\xcf\x11\xe0legacy")
    doc_id = _seed_doc(storage, src, "r.doc")
    with Session(engine) as db:
        res = recognize_standard_doc(db, storage, doc_id)
    assert res.recognition_status == "failed"
    assert res.error and (".doc" in res.error)
    assert _count_segments(doc_id) == 0


def test_recognize_doc_not_found_fails_without_raising(storage_dir, _test_db):
    storage = FileStorage(storage_dir)
    with Session(engine) as db:
        res = recognize_standard_doc(db, storage, 999999)
    assert res.recognition_status == "failed"
    assert res.segment_count == 0
    assert res.error  # 非空中文原因


def test_recognize_pdf_no_text_fails_with_scanned_hint(tmp_path, storage_dir, _test_db, monkeypatch):
    import app.recognition as recmod
    storage = FileStorage(storage_dir)
    src = tmp_path / "scan.pdf"
    src.write_bytes(b"%PDF-1.4 fake")  # 内容无所谓，parse_pdf 被 mock
    doc_id = _seed_doc(storage, src, "scan.pdf")
    monkeypatch.setattr(recmod, "parse_pdf", lambda path: [])
    with Session(engine) as db:
        res = recognize_standard_doc(db, storage, doc_id)
    assert res.recognition_status == "failed"
    assert res.segment_count == 0
    assert "扫描件" in (res.error or "")
    assert _count_segments(doc_id) == 0


import app.ocr as ocrmod  # noqa: E402


def _make_mixed_pdf(path: Path) -> None:
    """第1页有文本层(latin，pdfplumber 可抽)，第2页空白(模拟扫描页)。"""
    import fitz
    doc = fitz.open()
    p1 = doc.new_page()
    p1.insert_text((72, 72), "hello world")
    doc.new_page()  # 空白第2页
    doc.save(str(path))
    doc.close()


def test_parse_pdf_empty_page_falls_back_to_ocr(tmp_path, storage_dir, _test_db, monkeypatch):
    storage = FileStorage(storage_dir)
    src = tmp_path / "mixed.pdf"
    _make_mixed_pdf(src)
    doc_id = _seed_doc(storage, src, "mixed.pdf")
    monkeypatch.setattr(ocrmod, "OCR_AVAILABLE", True)
    monkeypatch.setattr(ocrmod, "ocr_pdf_page", lambda page: ["扫描出的文字"])
    with Session(engine) as db:
        res = recognize_standard_doc(db, storage, doc_id)
    assert res.recognition_status == "done"
    with Session(engine) as s:
        from sqlalchemy import text as sqltext
        rows = s.execute(sqltext(
            "SELECT content_text, page_no FROM parse_segment WHERE standard_doc_id=:d ORDER BY id"
        ), {"d": doc_id}).all()
    texts = [r[0] for r in rows]
    assert any("hello world" in t for t in texts)        # 文本页走 pdfplumber
    assert any("扫描出的文字" in t for t in texts)         # 空白页走 OCR
    assert any(r[1] == 2 for r in rows)                  # OCR 段 page_no=2


def test_recognize_image_file_via_ocr(tmp_path, storage_dir, _test_db, monkeypatch):
    storage = FileStorage(storage_dir)
    src = tmp_path / "scan.png"
    src.write_bytes(b"\x89PNG\r\n\x1a\n fake")  # 内容无所谓，ocr_image 被 mock
    doc_id = _seed_doc(storage, src, "scan.png")
    monkeypatch.setattr(ocrmod, "OCR_AVAILABLE", True)
    monkeypatch.setattr(ocrmod, "ocr_image", lambda path: ["图片里的中文"])
    with Session(engine) as db:
        res = recognize_standard_doc(db, storage, doc_id)
    assert res.recognition_status == "done"
    assert res.segment_count == 1
    with Session(engine) as s:
        from sqlalchemy import text as sqltext
        row = s.execute(sqltext(
            "SELECT content_text, segment_type, locator FROM parse_segment WHERE standard_doc_id=:d"
        ), {"d": doc_id}).first()
    assert "图片里的中文" in row[0]
    assert row[1] == "text"            # OCR 段落落 text


def test_recognize_image_without_ocr_degrades(tmp_path, storage_dir, _test_db, monkeypatch):
    storage = FileStorage(storage_dir)
    src = tmp_path / "scan.jpg"
    src.write_bytes(b"\xff\xd8\xff fake")
    doc_id = _seed_doc(storage, src, "scan.jpg")
    monkeypatch.setattr(ocrmod, "OCR_AVAILABLE", False)
    with Session(engine) as db:
        res = recognize_standard_doc(db, storage, doc_id)
    assert res.recognition_status == "failed"
    assert "OCR" in (res.error or "")
    assert _count_segments(doc_id) == 0


def test_reset_stuck_processing(tmp_path, storage_dir, _test_db):
    from app.recognition import reset_stuck_processing
    storage = FileStorage(storage_dir)
    src = tmp_path / "r.docx"
    from docx import Document
    d = Document(); d.add_paragraph("x"); d.save(str(src))
    doc_id = _seed_doc(storage, src, "r.docx")
    with Session(engine) as s:
        sd = s.get(StandardDoc, doc_id)
        sd.recognition_status = "processing"
        s.add(sd); s.commit()
    with Session(engine) as s:
        n = reset_stuck_processing(s)
    assert n == 1
    with Session(engine) as s:
        assert s.get(StandardDoc, doc_id).recognition_status == "pending"
