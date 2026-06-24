from io import BytesIO

from docx import Document


def _docx_bytes() -> bytes:
    d = Document()
    d.add_heading("第一章", level=1)
    d.add_paragraph("正文一段。")
    buf = BytesIO()
    d.save(buf)
    return buf.getvalue()


def test_upload_returns_processing_then_done(client):
    files = {"files": ("a.docx", _docx_bytes(),
                       "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    r = client.post("/api/standard-docs", files=files)
    assert r.status_code == 200
    up = r.json()["uploaded"][0]
    assert up["recognition_status"] == "processing"   # 立即返回 processing
    # TestClient 会在响应返回前跑完 BackgroundTasks，故列表此时已 done
    g = client.get("/api/standard-docs")
    doc = next(d for d in g.json() if d["doc_code"] == up["doc_code"])
    assert doc["recognition_status"] == "done"


def test_recognize_endpoint_returns_processing_then_done(client):
    files = {"files": ("b.docx", _docx_bytes(),
                       "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    up = client.post("/api/standard-docs", files=files).json()["uploaded"][0]
    doc_id = up["id"]
    r = client.post(f"/api/standard-docs/{doc_id}/recognize")
    assert r.status_code == 200
    assert r.json()["recognition_status"] == "processing"
    g = client.get("/api/standard-docs")
    doc = next(d for d in g.json() if d["id"] == doc_id)
    assert doc["recognition_status"] == "done"
