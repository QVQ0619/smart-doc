from app.config import get_max_upload_bytes
from app.main import app as fastapi_app


def test_upload_single_rule_file(client, storage_dir):
    r = client.post(
        "/api/standard-docs",
        files=[("files", ("政策A.pdf", b"rule-bytes", "application/pdf"))],
    )
    assert r.status_code == 200
    body = r.json()
    assert body["failed"] == []
    assert len(body["uploaded"]) == 1
    doc = body["uploaded"][0]
    assert doc["title"] == "政策A"
    assert doc["file_name"] == "政策A.pdf"
    assert doc["size_bytes"] == len(b"rule-bytes")
    assert doc["doc_code"].startswith("SD-")
    # 落到 standard_doc 子目录
    assert (storage_dir / "standard_doc").exists()
    assert list((storage_dir / "standard_doc").iterdir())


def test_upload_multiple_rule_files(client):
    r = client.post(
        "/api/standard-docs",
        files=[
            ("files", ("a.txt", b"aaa", "text/plain")),
            ("files", ("b.txt", b"bbbb", "text/plain")),
        ],
    )
    assert r.status_code == 200
    body = r.json()
    assert body["failed"] == []
    assert len(body["uploaded"]) == 2


def test_upload_oversize_goes_to_failed(client):
    fastapi_app.dependency_overrides[get_max_upload_bytes] = lambda: 3
    try:
        r = client.post(
            "/api/standard-docs",
            files=[
                ("files", ("ok.txt", b"ab", "text/plain")),
                ("files", ("big.txt", b"toolong", "text/plain")),
            ],
        )
    finally:
        fastapi_app.dependency_overrides.pop(get_max_upload_bytes, None)
    assert r.status_code == 200
    body = r.json()
    assert [d["file_name"] for d in body["uploaded"]] == ["ok.txt"]
    assert [f["name"] for f in body["failed"]] == ["big.txt"]


def test_standarddoc_failure_rolls_back_fileobject(client, storage_dir, monkeypatch):
    from app.routers import standard_docs
    from sqlalchemy import text
    from sqlmodel import Session
    from app.db import engine

    def boom(*a, **k):
        raise RuntimeError("forced standard_doc failure")
    monkeypatch.setattr(standard_docs, "StandardDoc", boom)

    # force=true 跳过同名检测(检测会用到被 mock 的 StandardDoc)，直达构造路径——
    # 本测试验证的是"StandardDoc 构造失败时回滚 file_object + 删盘"，与检测无关。
    r = client.post(
        "/api/standard-docs?force=true",
        files=[("files", ("x.txt", b"data", "text/plain"))],
    )
    assert r.status_code == 200
    body = r.json()
    assert body["uploaded"] == []
    assert len(body["failed"]) == 1
    # FileObject NOT orphaned (single-transaction rollback)
    with Session(engine) as s:
        n = s.execute(text("SELECT COUNT(*) FROM file_object")).scalar_one()
    assert n == 0
    # disk file removed
    sd_dir = storage_dir / "standard_doc"
    assert (not sd_dir.exists()) or list(sd_dir.iterdir()) == []


def _upload_one(client, name, content):
    return client.post(
        "/api/standard-docs",
        files=[("files", (name, content, "text/plain"))],
    ).json()["uploaded"][0]


def test_list_returns_active(client):
    _upload_one(client, "r1.txt", b"r1")
    _upload_one(client, "r2.txt", b"r2")
    rows = client.get("/api/standard-docs").json()
    assert {d["file_name"] for d in rows} == {"r1.txt", "r2.txt"}


def test_download_returns_original_bytes(client):
    doc = _upload_one(client, "p.txt", b"hello-bytes")
    r = client.get(f"/api/standard-docs/{doc['id']}/download")
    assert r.status_code == 200
    assert r.content == b"hello-bytes"


def test_download_unknown_404(client):
    assert client.get("/api/standard-docs/999999/download").status_code == 404


def test_delete_soft_removes_from_list_but_keeps_file(client, storage_dir):
    doc = _upload_one(client, "d.txt", b"data")
    assert any((storage_dir / "standard_doc").iterdir())

    r = client.delete(f"/api/standard-docs/{doc['id']}")
    assert r.status_code == 204
    # 列表已不含
    assert client.get("/api/standard-docs").json() == []
    # 磁盘文件保留（软删）
    assert list((storage_dir / "standard_doc").iterdir())
    # 再次下载 404（已软删）
    assert client.get(f"/api/standard-docs/{doc['id']}/download").status_code == 404


def test_delete_unknown_404(client):
    assert client.delete("/api/standard-docs/999999").status_code == 404


def test_download_404_when_file_object_soft_deleted_even_if_doc_active(client):
    # 与列表读路径口径一致: file_object.deleted_at 非空即不可下载,
    # 即便 standard_doc 仍 is_active (构造该中间态需直接改库)
    from datetime import datetime

    from sqlmodel import Session

    from app.db import engine
    from app.models import FileObject, StandardDoc

    doc = _upload_one(client, "x.txt", b"bytes")
    with Session(engine) as s:
        sd = s.get(StandardDoc, doc["id"])
        fo = s.get(FileObject, sd.file_id)
        fo.deleted_at = datetime.now()
        s.add(fo)
        s.commit()

    assert client.get(f"/api/standard-docs/{doc['id']}/download").status_code == 404


def test_upload_docx_auto_recognized(client):
    import io
    from docx import Document
    d = Document()
    d.add_heading("第一章", level=1)
    d.add_paragraph("正文。")
    buf = io.BytesIO()
    d.save(buf)
    r = client.post(
        "/api/standard-docs",
        files=[("files", ("规范.docx", buf.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))],
    )
    assert r.status_code == 200
    doc = r.json()["uploaded"][0]
    assert doc["recognition_status"] == "done"
    assert doc["segment_count"] >= 2


def test_list_includes_recognition_status(client):
    client.post("/api/standard-docs", files=[("files", ("a.txt", b"x", "text/plain"))])
    rows = client.get("/api/standard-docs").json()
    assert "recognition_status" in rows[0]
    # .txt 不支持 → failed
    assert rows[0]["recognition_status"] == "failed"


def test_recognize_endpoint_returns_result(client):
    import io
    from docx import Document
    d = Document()
    d.add_paragraph("仅一段。")
    buf = io.BytesIO()
    d.save(buf)
    doc = client.post(
        "/api/standard-docs",
        files=[("files", ("b.docx", buf.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))],
    ).json()["uploaded"][0]
    r = client.post(f"/api/standard-docs/{doc['id']}/recognize")
    assert r.status_code == 200
    body = r.json()
    assert body["doc_id"] == doc["id"]
    assert body["recognition_status"] == "done"
    assert body["segment_count"] >= 1


def test_recognize_unknown_404(client):
    assert client.post("/api/standard-docs/999999/recognize").status_code == 404
