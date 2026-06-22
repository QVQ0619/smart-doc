from pathlib import Path

from docx import Document

from app.recognition import parse_pdf, parse_docx, SegmentDraft

FIXTURE = Path(__file__).parent / "fixtures" / "sample.pdf"


def _make_docx(path: Path) -> None:
    d = Document()
    d.add_heading("第一章 总则", level=1)
    d.add_paragraph("本办法适用于立项审查。")
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "项目"
    t.cell(0, 1).text = "金额"
    t.cell(1, 0).text = "A"
    t.cell(1, 1).text = "100"
    d.save(str(path))


def test_parse_docx_emits_heading_paragraph_table(tmp_path):
    p = tmp_path / "rule.docx"
    _make_docx(p)
    segs = parse_docx(p)
    kinds = [s.segment_type for s in segs]
    assert "heading" in kinds and "paragraph" in kinds and "table" in kinds
    heading = next(s for s in segs if s.segment_type == "heading")
    assert heading.content_text == "第一章 总则"
    assert heading.page_no is None
    assert heading.locator.get("level") == 1
    table = next(s for s in segs if s.segment_type == "table")
    assert "项目 | 金额" in table.content_text
    assert all(s.page_no is None for s in segs)


def test_parse_pdf_two_pages_with_table():
    segs = parse_pdf(FIXTURE)
    assert {s.page_no for s in segs} == {1, 2}
    texts = " ".join(s.content_text for s in segs)
    assert "FIRST PAGE PARAGRAPH ONE" in texts
    assert "SECOND PAGE PARAGRAPH" in texts
    p1 = [s for s in segs if s.page_no == 1 and s.segment_type == "paragraph"]
    assert p1 and p1[0].locator == {"page": 1, "block_index": 0}
    tables = [s for s in segs if s.segment_type == "table"]
    assert tables and "CELL_1" in tables[0].content_text
    p2_paras = [s for s in segs if s.page_no == 2 and s.segment_type == "paragraph"]
    assert all("CELL_1" not in s.content_text for s in p2_paras)
    assert all("COL_A" not in s.content_text for s in p2_paras)


def test_segmentdraft_is_dataclass():
    s = SegmentDraft(page_no=3, locator={"page": 3}, segment_type="paragraph", content_text="x")
    assert s.page_no == 3 and s.locator == {"page": 3}
